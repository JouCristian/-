# -*- coding: utf-8 -*-
"""
CSP Word/PDF 题解复盘生成器 v9
- 更紧凑的自然段排版：自动合并过碎短行
- 更清晰的轻量彩色卡片/表格/图块
- C++ 代码块高亮
- 自动目录
- Word 后自动尝试生成 PDF

依赖：pip install python-docx
可选 PDF：pip install docx2pdf
运行：python csp_word_generator_v9.py input.txt
"""
from __future__ import annotations
import re, sys, shutil, subprocess
from pathlib import Path
from typing import Dict, List, Tuple
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

FONT_CN = "黑体"
FONT_CODE = "Consolas"
DARK = "1F2937"
BLUE = "1F4E79"
BLUE_LIGHT = "D9EAF7"
GRAY = "F3F4F6"
GREEN = "EAF4EA"
RED = "FDECEC"

CPP_KEYWORDS = set("""int long double float char bool void return for while if else switch case break continue using namespace std vector string include cin cout endl const auto struct class public private true false push_back size swap""".split())


def read_text(path: Path) -> str:
    for enc in ("utf-8-sig", "utf-8", "gb18030"):
        try:
            return path.read_text(encoding=enc)
        except UnicodeDecodeError:
            pass
    raise RuntimeError("无法识别 input.txt 编码，请保存为 UTF-8。")


def parse_sections(text: str) -> Dict[str, str]:
    pat = re.compile(r"^【(.+?)】\s*$", re.M)
    ms = list(pat.finditer(text))
    ans = {}
    for i, m in enumerate(ms):
        key = m.group(1).strip()
        start = m.end()
        end = ms[i + 1].start() if i + 1 < len(ms) else len(text)
        ans[key] = text[start:end].strip()
    return ans


def clean_filename(title: str) -> str:
    s = title.replace("CSP真题：", "").replace("——完整题解复盘（C++）", "")
    s = re.sub(r"[\\/:*?\"<>|]", "_", s)
    s = re.sub(r"\s+", "", s).strip("_-")
    return s or "未命名题目"


def font(run, name=FONT_CN, size=10.5, bold=False, color=DARK):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd"); tcPr.append(shd)
    shd.set(qn("w:fill"), fill)


def border(cell, color="D1D5DB", sz="6"):
    tcPr = cell._tc.get_or_add_tcPr()
    b = tcPr.first_child_found_in("w:tcBorders")
    if b is None:
        b = OxmlElement("w:tcBorders"); tcPr.append(b)
    for e in ("top", "left", "bottom", "right"):
        x = b.find(qn("w:" + e))
        if x is None:
            x = OxmlElement("w:" + e); b.append(x)
        x.set(qn("w:val"), "single"); x.set(qn("w:sz"), sz); x.set(qn("w:color"), color)


def spacing(p, before=0, after=3, line=1.12):
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line


def add_text_para(doc, text, size=10.5, color=DARK, before=0, after=3):
    p = doc.add_paragraph(); spacing(p, before, after)
    r = p.add_run(text); font(r, size=size, color=color)
    return p


def heading(doc, text, level=1):
    p = doc.add_paragraph(); spacing(p, before=9 if level == 1 else 5, after=4)
    r = p.add_run(text); font(r, size=15 if level == 1 else 12, bold=True, color=BLUE)
    if level == 1:
        pPr = p._p.get_or_add_pPr(); pBdr = OxmlElement("w:pBdr"); bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "10"); bottom.set(qn("w:space"), "3"); bottom.set(qn("w:color"), "8BB7D8")
        pBdr.append(bottom); pPr.append(pBdr)


def setup(doc):
    sec = doc.sections[0]
    sec.top_margin = Cm(1.8); sec.bottom_margin = Cm(1.8); sec.left_margin = Cm(2.0); sec.right_margin = Cm(2.0)
    st = doc.styles["Normal"]
    st.font.name = FONT_CN; st._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN); st.font.size = Pt(10.5)


def add_title(doc, title, subtitle):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; spacing(p, before=6, after=5)
    r = p.add_run(title); font(r, size=20, bold=True, color=BLUE)
    if subtitle:
        p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER; spacing(p2, after=10)
        r2 = p2.add_run(subtitle); font(r2, size=10.5, color="6B7280")


def add_toc(doc, heads):
    heading(doc, "目录", 1)
    table = doc.add_table(rows=0, cols=2); table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(heads, 1):
        row = table.add_row(); row.cells[0].text = str(i); row.cells[1].text = h
        for c in row.cells:
            border(c, "FFFFFF", "0")
            for p in c.paragraphs:
                spacing(p, after=1)
                for r in p.runs: font(r, size=10)
    doc.add_paragraph()


def compact(text: str) -> str:
    lines = [x.rstrip() for x in text.splitlines()]
    out, buf = [], []
    in_code = in_fig = False
    def special(s):
        st = s.strip()
        if not st or st.startswith(("```", "[[图:")) or st == "[[/图]]": return True
        if st.startswith("|") and st.endswith("|"): return True
        if st.startswith("•"): return True
        if re.match(r"^阶段\s*\d+[：:]", st): return True
        if re.match(r"^[\d\s\-]+$", st) and any(ch.isdigit() for ch in st): return True
        if re.match(r"^(输入|输出|解释|原矩阵|转置后|展开后|查询|结果)[：:]?$", st): return True
        return False
    def flush():
        nonlocal buf
        if buf:
            s = "".join(buf).strip()
            if s: out.append(s)
            buf = []
    for raw in lines:
        st = raw.strip()
        if st.startswith("```"):
            flush(); out.append(raw); in_code = not in_code; continue
        if in_code: out.append(raw); continue
        if st.startswith("[[图:"):
            flush(); out.append(raw); in_fig = True; continue
        if st == "[[/图]]":
            flush(); out.append(raw); in_fig = False; continue
        if in_fig: out.append(raw); continue
        if not st:
            flush();
            if out and out[-1] != "": out.append("")
            continue
        if special(st):
            flush(); out.append(raw)
        else:
            buf.append(st)
    flush()
    final, prev = [], False
    for x in out:
        blank = not x.strip()
        if blank and prev: continue
        final.append(x); prev = blank
    return "\n".join(final).strip()


def md_table(lines, i):
    return i + 1 < len(lines) and lines[i].strip().startswith("|") and re.match(r"^\|\s*[-: ]+\|", lines[i+1].strip())


def add_table(doc, rows):
    if not rows: return
    table = doc.add_table(rows=len(rows), cols=max(len(r) for r in rows)); table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row in enumerate(rows):
        for j in range(len(table.columns)):
            c = table.cell(i, j); c.text = row[j] if j < len(row) else ""
            shade(c, BLUE_LIGHT if i == 0 else "FFFFFF"); border(c)
            for p in c.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER; spacing(p, after=1)
                for r in p.runs: font(r, size=9.5, bold=(i == 0))
    doc.add_paragraph()


def add_code(doc, code):
    table = doc.add_table(rows=1, cols=1); table.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = table.cell(0,0); c.text = ""; shade(c, GRAY); border(c)
    for line in code.rstrip().splitlines():
        p = c.add_paragraph(); spacing(p, after=0, line=1.0)
        for tok in re.split(r"(\W+)", line):
            if tok == "": continue
            r = p.add_run(tok)
            if tok in CPP_KEYWORDS: font(r, FONT_CODE, 9, True, "0F4C81")
            elif re.fullmatch(r"\d+", tok): font(r, FONT_CODE, 9, False, "9A3412")
            else: font(r, FONT_CODE, 9, False, DARK)
    try: c._element.remove(c.paragraphs[0]._element)
    except Exception: pass
    doc.add_paragraph()


def matrix(s): return "\n".join(x.strip() for x in s.split(";") if x.strip())

def parse_fig(lines, i):
    data = {}; i += 1
    while i < len(lines):
        st = lines[i].strip()
        if st == "[[/图]]": return data, i + 1
        if "：" in st:
            k,v = st.split("：",1); data[k.strip()] = v.strip()
        elif ":" in st:
            k,v = st.split(":",1); data[k.strip()] = v.strip()
        i += 1
    return data, i


def add_fig(doc, kind, data):
    p = doc.add_paragraph(); spacing(p, before=5, after=2)
    r = p.add_run("图：" + data.get("标题", kind)); font(r, size=10.5, bold=True, color=BLUE)
    if data.get("说明"): add_text_para(doc, data["说明"], size=9.5, color="4B5563", after=2)
    if kind == "reshape":
        table = doc.add_table(rows=2, cols=5); table.alignment = WD_TABLE_ALIGNMENT.CENTER
        vals = [(data.get("左标题","原矩阵"), matrix(data.get("左矩阵",""))), ("→","→"), (data.get("中标题","序列"), data.get("一维序列","")), ("→","→"), (data.get("右标题","结果"), matrix(data.get("右矩阵","")))]
        for j,(h,v) in enumerate(vals): table.cell(0,j).text = "" if h == "→" else h; table.cell(1,j).text = v
    elif kind == "transpose":
        table = doc.add_table(rows=2, cols=3); table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.cell(0,0).text=data.get("左标题","原矩阵"); table.cell(0,1).text=""; table.cell(0,2).text=data.get("右标题","转置后")
        table.cell(1,0).text=matrix(data.get("左矩阵","")); table.cell(1,1).text="→"; table.cell(1,2).text=matrix(data.get("右矩阵",""))
    else:
        table = doc.add_table(rows=3, cols=2); table.alignment = WD_TABLE_ALIGNMENT.CENTER
        pairs=[(data.get("左标题","矩阵"),matrix(data.get("矩阵",""))),(data.get("右标题","下标"),data.get("下标","")),("公式",data.get("公式",""))]
        for i,(k,v) in enumerate(pairs): table.cell(i,0).text=k; table.cell(i,1).text=v
    for i,row in enumerate(table.rows):
        for c in row.cells:
            shade(c, BLUE_LIGHT if i==0 else "FFFFFF"); border(c,"CBD5E1"); c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in c.paragraphs:
                p.alignment=WD_ALIGN_PARAGRAPH.CENTER; spacing(p, after=1, line=1.05)
                for r in p.runs: font(r, size=9.5)
    doc.add_paragraph()


def add_box(doc, title, text, fill):
    table=doc.add_table(rows=1, cols=1); table.alignment=WD_TABLE_ALIGNMENT.CENTER
    c=table.cell(0,0); c.text=""; shade(c, fill); border(c)
    p=c.paragraphs[0]; spacing(p, before=3, after=3)
    r=p.add_run(title+"："); font(r, size=10.5, bold=True, color=BLUE)
    r=p.add_run(compact(text).replace("\n"," ")); font(r, size=10.5)
    doc.add_paragraph()


def add_sample(doc, label, text):
    if not text.strip(): return
    p=doc.add_paragraph(); spacing(p, before=3, after=1); r=p.add_run(label); font(r, size=10.5, bold=True, color=BLUE)
    table=doc.add_table(rows=1, cols=1); table.alignment=WD_TABLE_ALIGNMENT.CENTER
    c=table.cell(0,0); c.text=""; shade(c, GRAY); border(c)
    for line in text.strip().splitlines():
        p2=c.add_paragraph(); spacing(p2, after=0, line=1.0); r=p2.add_run(line); font(r, FONT_CODE, 9.5)
    try: c._element.remove(c.paragraphs[0]._element)
    except Exception: pass


def rich(doc, text):
    lines=compact(text).splitlines(); i=0; in_code=False; buf=[]
    while i < len(lines):
        st=lines[i].strip()
        if st.startswith("```"):
            if not in_code: in_code=True; buf=[]
            else: in_code=False; add_code(doc,"\n".join(buf))
            i+=1; continue
        if in_code: buf.append(lines[i]); i+=1; continue
        if st.startswith("[[图:"):
            kind=st[4:-2].strip(); data,nxt=parse_fig(lines,i); add_fig(doc,kind,data); i=nxt; continue
        if not st: i+=1; continue
        if md_table(lines,i):
            rows=[]
            while i < len(lines) and lines[i].strip().startswith("|"):
                parts=[x.strip() for x in lines[i].strip().strip("|").split("|")]
                if not all(re.fullmatch(r"[-: ]+", x) for x in parts): rows.append(parts)
                i+=1
            add_table(doc, rows); continue
        if st.startswith("•"):
            p=doc.add_paragraph(); p.paragraph_format.left_indent=Cm(0.45); spacing(p, after=2)
            r=p.add_run(st); font(r, size=10.5); i+=1; continue
        if re.match(r"^阶段\s*\d+[：:]", st):
            p=doc.add_paragraph(); spacing(p, before=5, after=2); r=p.add_run(st); font(r, size=11, bold=True, color=BLUE); i+=1; continue
        add_text_para(doc, st, after=3); i+=1
    if in_code and buf: add_code(doc,"\n".join(buf))


def keywords(doc, text):
    kws=[x.strip() for x in re.split(r"[；;、,\n]+", text) if x.strip()]
    table=doc.add_table(rows=(len(kws)+3)//4, cols=4); table.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,row in enumerate(table.rows):
        for j,c in enumerate(row.cells):
            idx=i*4+j; c.text=kws[idx] if idx < len(kws) else ""; shade(c,"F8FAFC" if idx < len(kws) else "FFFFFF"); border(c)
            for p in c.paragraphs:
                p.alignment=WD_ALIGN_PARAGRAPH.CENTER; spacing(p, after=1)
                for r in p.runs: font(r, size=9.5)
    doc.add_paragraph()


def pdf(docx):
    pdfp=docx.with_suffix(".pdf")
    try:
        from docx2pdf import convert
        convert(str(docx), str(pdfp)); print("PDF 已生成：", pdfp); return
    except Exception as e:
        print("docx2pdf 不可用，尝试 LibreOffice：", e)
    soffice=shutil.which("soffice") or shutil.which("libreoffice")
    if soffice:
        try:
            subprocess.run([soffice,"--headless","--convert-to","pdf","--outdir",str(docx.parent),str(docx)],check=True)
            print("PDF 已生成：", pdfp)
        except Exception as e: print("PDF 生成失败：", e)
    else: print("未找到 Word/docx2pdf 或 LibreOffice，已跳过 PDF。")


def build(input_path: Path):
    sec=parse_sections(read_text(input_path))
    title=sec.get("标题","CSP真题：未命名——完整题解复盘（C++）")
    subtitle=sec.get("副标题","")
    out=input_path.parent / f"CSP_{clean_filename(title)}_完整题解复盘.docx"
    doc=Document(); setup(doc); add_title(doc,title,subtitle)
    order=["题目背景","完整题目要求","输入格式","输出格式","输入输出样例","子任务与限制条件","题目提示与关键区别","我的完整思考流程复盘","原始代码","错误分析","优化过程","最终AC代码","可进一步优化的小细节","复杂度分析","题解关键词","总结"]
    heads=[h for h in order if h=="输入输出样例" or sec.get(h,"").strip()]
    add_toc(doc,heads)
    cn="一二三四五六七八九十"; idx=1
    def numbered(h):
        nonlocal idx
        pre=cn[idx-1] if idx<=10 else str(idx); heading(doc,f"{pre}、{h}"); idx+=1
    for h in ["题目背景","完整题目要求","输入格式","输出格式"]:
        if sec.get(h,"").strip(): numbered(h); rich(doc,sec[h])
    if any(sec.get(f"样例{n}输入","").strip() for n in (1,2)):
        numbered("输入输出样例")
        for n in (1,2):
            if any(sec.get(f"样例{n}{x}","").strip() for x in ("输入","输出","解释")):
                heading(doc,f"样例 {n}",2); add_sample(doc,"输入：",sec.get(f"样例{n}输入","")); add_sample(doc,"输出：",sec.get(f"样例{n}输出",""))
                if sec.get(f"样例{n}解释","").strip():
                    p=doc.add_paragraph(); spacing(p,before=3,after=1); r=p.add_run("解释："); font(r,size=10.5,bold=True,color=BLUE)
                    rich(doc,sec[f"样例{n}解释"])
    for h in ["子任务与限制条件","题目提示与关键区别","我的完整思考流程复盘","原始代码","错误分析","优化过程","最终AC代码","可进一步优化的小细节","复杂度分析","题解关键词","总结"]:
        if not sec.get(h,"").strip(): continue
        numbered(h)
        if h=="题解关键词": keywords(doc,sec[h])
        elif h=="错误分析": add_box(doc,h,sec[h],RED)
        elif h=="优化过程": add_box(doc,h,sec[h],GREEN)
        else: rich(doc,sec[h])
    footer=doc.sections[0].footer.paragraphs[0]; footer.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=footer.add_run("CSP 题解复盘"); font(r,size=9,color="9CA3AF")
    doc.save(out); print("Word 已生成：", out); pdf(out)

if __name__ == "__main__":
    path=Path(sys.argv[1]) if len(sys.argv)>1 else Path("input.txt")
    build(path)
